"""
MPR Markdown Parser — agents/mpr_parser.py
Parses the structured Monthly Progress Report .md file (EL/02 template)
into a typed exec_data dict ready for the compliance engine + risk predictor.

11 sections parsed:
  S1  — Project metadata
  S2  — Physical & financial progress
  S3  — BoQ execution table
  S4  — Material reconciliation table
  S5  — Labour & machinery deployment
  S6  — QA results + NCRs/RFIs
  S7  — External disruptions & rainfall
  S8  — Land acquisition & utilities
  S9  — GFC drawing status
  S10 — RA bill & payment
  S11 — Site engineer declaration
"""
import re
from datetime import date, datetime
from typing import Any, Optional


# ── Helpers ──────────────────────────────────────────────────────────────

def _kv(text: str, label: str, cast=str) -> Optional[Any]:
    """Extract a key-value field from MPR markdown: '- **Label (...):** value' OR '| Label | value |'"""
    # Try list pattern first
    pattern = rf"\*\*{re.escape(label)}(?:\s*\([^)]*\))?\s*[:\*]+\**\s*(.+)"
    m = re.search(pattern, text, re.IGNORECASE)
    raw = None
    if m:
        raw = m.group(1).strip().rstrip("*").strip()
    else:
        # Try table pattern: | Label | value |
        # Optional parenthetical after label e.g. | Label (km) | value |
        pattern_tbl = rf"\|\s*{re.escape(label)}(?:\s*\([^)]*\))?\s*\|\s*(.+?)\s*\|"
        m_tbl = re.search(pattern_tbl, text, re.IGNORECASE)
        if m_tbl:
            raw = m_tbl.group(1).strip()
            
    if not raw or raw in ("[OPTIONAL]", "[Name]", "", "-"):
        return None
    try:
        return cast(raw.replace(",", "").replace("₹", "").replace("Rs.", "").strip())
    except (ValueError, TypeError):
        return None


def _table(text: str, section_header: str) -> list[dict]:
    """Parse a markdown table under a given section header. Returns list of row dicts."""
    # Find the section (make 'Section' optional to support various markdown formats)
    section_pattern = rf"##\s*(?:Section\s*[\dA-Za-z\.\-]*\s*[—\-:]?\s*)?{re.escape(section_header)}.*?\n(.*?)(?=\n##|\Z)"
    sm = re.search(section_pattern, text, re.IGNORECASE | re.DOTALL)
    if not sm:
        return []

    block = sm.group(1)
    lines = [l.strip() for l in block.splitlines() if l.strip().startswith("|")]
    if len(lines) < 2:
        return []

    headers = [h.strip() for h in lines[0].split("|") if h.strip()]
    rows = []
    for line in lines[2:]:  # skip separator line
        cells = [c.strip() for c in line.split("|") if c.strip() != ""]
        if len(cells) == len(headers):
            rows.append(dict(zip(headers, cells)))
    return rows




# ── Validation ────────────────────────────────────────────────────────────

class MPRValidationError(Exception):
    def __init__(self, errors: list):
        self.errors = errors
        super().__init__(f"MPR validation failed: {errors}")


def validate_mpr(record: dict, prev_actual_pct: float = 0.0, bypass_date_check: bool = False) -> list:
    """
    Run all 5 validation rules from EL/02.
    Returns list of error strings (empty = pass).
    bypass_date_check: set True for demo/synthetic documents dated in the future.
    """
    errors = []

    # Rule 1 — Date check: period end ≤ today
    period_end = record.get("reporting_period_end")
    if period_end and not bypass_date_check:
        try:
            end_date = datetime.strptime(period_end, "%Y-%m-%d").date()
            if end_date > date.today():
                errors.append(f"Reporting period end ({period_end}) is in the future.")
        except ValueError:
            errors.append(f"Cannot parse reporting_period_end: '{period_end}'")

    # Rule 2 — Monotonicity: actual_pct ≥ previous month
    actual = record.get("actual_physical_pct", 0)
    if actual < prev_actual_pct:
        errors.append(
            f"actual_physical_pct ({actual}%) is less than previous month ({prev_actual_pct}%). "
            f"Physical progress cannot decrease."
        )

    # Rule 3 — Labour sanity: actual ≤ 150% of planned
    skilled_planned = record.get("labour_skilled_planned", 0)
    skilled_actual = record.get("labour_skilled_actual", 0)
    if skilled_planned > 0 and skilled_actual > skilled_planned * 1.5:
        errors.append(
            f"labour_skilled_actual ({skilled_actual}) exceeds 150% of planned ({skilled_planned}). "
            f"Likely data entry error."
        )

    # Rule 4 — QA consistency: passed + failed = conducted
    for row in record.get("qa_results", []):
        conducted = _safe_int(row.get("Tests Conducted") or row.get("tests_conducted"))
        passed = _safe_int(row.get("Tests Passed") or row.get("tests_passed"))
        failed = _safe_int(row.get("Tests Failed") or row.get("tests_failed"))
        material = row.get("Material", "")
        if conducted > 0 and (passed + failed) != conducted:
            errors.append(
                f"QA inconsistency for '{material}': "
                f"passed({passed}) + failed({failed}) ≠ conducted({conducted})"
            )

    # Rule 5 — RA bill date after reporting period
    ra_date = record.get("ra_bill_submitted_date")
    if period_end and ra_date and ra_date < period_end:
        errors.append(
            f"RA bill submitted ({ra_date}) before reporting period ends ({period_end})."
        )

    return errors


# ── Main Parser ───────────────────────────────────────────────────────────

def extract_from_table_row(tables, row_label: str, col_idx: int) -> str:
    """Finds a row starting with row_label in any table and returns the text at col_idx."""
    for table in tables:
        for row in table.rows:
            if not row.cells: continue
            if row_label.lower() in row.cells[0].text.lower():
                if col_idx < len(row.cells):
                    return row.cells[col_idx].text.strip()
    return ""

def _safe_float(val) -> float:
    if val is None: return 0.0
    s = str(val).split(' ')[0] # handle '214 of 730' -> 214
    m = re.search(r"[-+]?\d*\.\d+|\d+", s.replace(",", ""))
    if m:
        return float(m.group())
    return 0.0

def _safe_int(val) -> int:
    return int(_safe_float(val))

def parse_mpr_docx(file_path_or_bytes, prev_actual_pct: float = 0.0, bypass_date_check: bool = False) -> dict:
    import docx
    doc = docx.Document(file_path_or_bytes)
    
    # Extract text from paragraphs
    text = "\n".join([p.text for p in doc.paragraphs])
    
    # ── Section 1 — Project metadata ────────────────────────────────────
    project_name   = extract_from_table_row(doc.tables, "Project Name", 1) or _kv(text, "Project Name")
    agreement_no   = extract_from_table_row(doc.tables, "Agreement Number", 1) or _kv(text, "Agreement Number")
    contractor     = extract_from_table_row(doc.tables, "Contractor", 1) or _kv(text, "Contractor Name")
    eic            = extract_from_table_row(doc.tables, "Engineer-in-Charge", 1) or _kv(text, "Engineer-in-Charge")
    period_raw     = extract_from_table_row(doc.tables, "Reporting Period", 1) or _kv(text, "Reporting Period")
    scd_raw        = extract_from_table_row(doc.tables, "Scheduled Completion", 1) or _kv(text, "Stipulated Date of Completion")
    eot_date       = extract_from_table_row(doc.tables, "Approved EoT", 1) or _kv(text, "Approved EoT Date")
    day_number_raw = extract_from_table_row(doc.tables, "Day Number", 1) or _kv(text, "Day Number")
    
    day_number = _safe_int(day_number_raw)

    period_end = None
    if period_raw and " to " in period_raw:
        period_end = period_raw.split(" to ")[1].strip()
        # Normalise to YYYY-MM-DD
        for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                period_end = datetime.strptime(period_end, fmt).strftime("%Y-%m-%d")
                break
            except ValueError:
                continue

    # ── Section 2 — Physical & financial progress ───────────────────────
    planned_physical_pct    = _safe_float(extract_from_table_row(doc.tables, "Physical Progress", 1) or _kv(text, "Planned Physical Progress to Date"))
    actual_physical_pct     = _safe_float(extract_from_table_row(doc.tables, "Physical Progress", 2) or _kv(text, "Actual Physical Progress to Date"))
    variance_pct            = _safe_float(extract_from_table_row(doc.tables, "Physical Progress", 3) or _kv(text, "Variance"))
    cumulative_exp_inr      = _safe_float(extract_from_table_row(doc.tables, "Cumulative Expenditure", 2) or _kv(text, "Cumulative Expenditure to Date"))
    planned_exp_inr         = _safe_float(extract_from_table_row(doc.tables, "Cumulative Expenditure", 1) or _kv(text, "Planned Expenditure to Date"))
    financial_progress_pct  = _safe_float(extract_from_table_row(doc.tables, "Financial Progress", 2) or _kv(text, "Financial Progress"))

    # ── Section 3 — BoQ table ────────────────────────────────────────────
    boq_items = []
    # ── Section 4 — Material reconciliation ─────────────────────────────
    materials = []
    # ── Section 5 — Labour & machinery ───────────────────────────────────
    labour_skilled_planned   = _safe_int(extract_from_table_row(doc.tables, "Planned Skilled Labour", 1) or _kv(text, "Planned Skilled Labour"))
    labour_skilled_actual    = _safe_int(extract_from_table_row(doc.tables, "Actual Skilled Labour", 1) or _kv(text, "Actual Skilled Labour"))
    labour_unskilled_planned = _safe_int(extract_from_table_row(doc.tables, "Planned Unskilled Labour", 1) or _kv(text, "Planned Unskilled Labour"))
    labour_unskilled_actual  = _safe_int(extract_from_table_row(doc.tables, "Actual Unskilled Labour", 1) or _kv(text, "Actual Unskilled Labour"))
    machinery_idle_days      = _safe_int(extract_from_table_row(doc.tables, "Machinery Idle", 1) or _kv(text, "Machinery Idle Days"))
    machinery_deploy_raw     = extract_from_table_row(doc.tables, "Machinery Deployment", 1) or extract_from_table_row(doc.tables, "Machinery Deployed", 1) or _kv(text, "Machinery Deployment")
    machinery_deployment_pct = _safe_float(machinery_deploy_raw) if machinery_deploy_raw else 80.0

    skilled_util = ((labour_skilled_actual / labour_skilled_planned * 100) if labour_skilled_planned else 100.0)
    unskilled_util = ((labour_unskilled_actual / labour_unskilled_planned * 100) if labour_unskilled_planned else 100.0)

    # ── Section 6 — QA ────────────────────────────────────────────────────
    qa_results = []
    total_conducted = 0
    total_failed = 0
    for table in doc.tables:
        if not table.rows: continue
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if 'conducted' in headers or 'tests conducted' in headers:
            idx_cond = headers.index('conducted') if 'conducted' in headers else headers.index('tests conducted')
            idx_fail = headers.index('failed') if 'failed' in headers else (headers.index('tests failed') if 'tests failed' in headers else -1)
            for row in table.rows[1:]:
                if len(row.cells) > idx_cond:
                    total_conducted += _safe_int(row.cells[idx_cond].text)
                if idx_fail >= 0 and len(row.cells) > idx_fail:
                    total_failed += _safe_int(row.cells[idx_fail].text)
            break
            
    test_fail_rate = (total_failed / total_conducted * 100) if total_conducted > 0 else 0.0

    ncrs_issued  = _safe_int(extract_from_table_row(doc.tables, "NCRs Issued", 1) or _kv(text, "NCRs Issued This Month"))
    ncrs_pending = _safe_int(extract_from_table_row(doc.tables, "NCRs Pending", 1) or _kv(text, "NCRs Pending Closure"))
    rfis_submitted = _safe_int(extract_from_table_row(doc.tables, "RFIs Submitted", 1) or _kv(text, "RFIs Submitted"))
    rfis_approved  = _safe_int(extract_from_table_row(doc.tables, "RFIs Approved", 1) or _kv(text, "RFIs Approved"))
    rfis_pending   = _safe_int(extract_from_table_row(doc.tables, "RFIs Pending", 1) or _kv(text, "RFIs Pending"))
    
    # ── Section 7 — Disruptions & weather ────────────────────────────────
    working_days           = _safe_int(extract_from_table_row(doc.tables, "Working Days", 1) or _kv(text, "Working Days in Month") or 31)
    days_lost_rainfall     = _safe_int(extract_from_table_row(doc.tables, "Days Lost to Rainfall", 1) or _kv(text, "Days Lost to Rainfall"))
    days_lost_other        = _safe_int(extract_from_table_row(doc.tables, "Days Lost to Other", 1) or _kv(text, "Days Lost to Other Hindrances"))
    daily_rainfall_mm      = _safe_float(extract_from_table_row(doc.tables, "Daily Rainfall", 1) or _kv(text, "Daily Rainfall"))
    cumulative_rainfall_mm = _safe_float(extract_from_table_row(doc.tables, "Cumulative Rainfall", 1) or _kv(text, "Cumulative Rainfall this Month"))

    # ── Section 8 — Land & utilities ─────────────────────────────────────
    row_total_km   = _safe_float(extract_from_table_row(doc.tables, "Total RoW Required", 1) or _kv(text, "Total RoW Required") or 40.0)
    row_handover_km = _safe_float(extract_from_table_row(doc.tables, "RoW Handed Over", 1) or _kv(text, "RoW Handed Over") or row_total_km)
    row_pending_km  = _safe_float(extract_from_table_row(doc.tables, "RoW Pending", 1) or _kv(text, "RoW Pending"))
    row_handover_pct = (row_handover_km / row_total_km * 100) if row_total_km > 0 else 100.0

    utility_raw = extract_from_table_row(doc.tables, "Utility Shifting Status", 1) or _kv(text, "Utility Shifting Status") or ""
    utility_shifting_pending = bool(utility_raw and "pending" in utility_raw.lower())
    tree_felling_raw = extract_from_table_row(doc.tables, "Tree Felling", 1) or _kv(text, "Tree Felling Clearance") or "Y"
    forest_clearance_pending = tree_felling_raw.strip().upper() == "N"

    # ── Section 9 — GFC drawings ─────────────────────────────────────────
    gfc_total    = _safe_int(extract_from_table_row(doc.tables, "Total GFC Drawings Required", 1) or _kv(text, "Total GFC Drawings Required"))
    gfc_approved = _safe_int(extract_from_table_row(doc.tables, "GFC Drawings Approved", 1) or _kv(text, "GFC Drawings Approved"))
    gfc_pending  = _safe_int(extract_from_table_row(doc.tables, "GFC Drawings Pending", 1) or _kv(text, "GFC Drawings Pending"))

    # ── Section 10 — RA bill & payment ───────────────────────────────────
    ra_bill_no         = extract_from_table_row(doc.tables, "RA Bill Number", 1) or _kv(text, "RA Bill Number")
    ra_bill_amount_inr = _safe_float(extract_from_table_row(doc.tables, "RA Bill Amount", 1) or _kv(text, "RA Bill Amount"))
    ra_bill_submitted  = extract_from_table_row(doc.tables, "RA Bill Submitted", 1) or _kv(text, "RA Bill Submitted Date") or period_end
    prev_payment_yn    = extract_from_table_row(doc.tables, "Previous Bill Payment Received", 1) or _kv(text, "Previous Bill Payment Received") or "Y"
    prev_payment_date  = extract_from_table_row(doc.tables, "Previous Bill Payment Date", 1) or _kv(text, "Previous Bill Payment Date")
    payment_delay_days = _safe_int(extract_from_table_row(doc.tables, "Payment Delay", 1) or _kv(text, "Payment Delay"))
    payment_delayed = payment_delay_days > 0

    ld_deducted = _safe_float(
        extract_from_table_row(doc.tables, "Cumulative LD Deducted", 1)
        or extract_from_table_row(doc.tables, "LD Deducted", 1)
        or _kv(text, "Cumulative LD Deducted")
    )
    if ld_deducted > 0 and ld_deducted < 10000:
        ld_deducted = ld_deducted * 100000

    open_ncrs = [
        {"id": f"NCR-{i+1:03d}", "issued_date": period_end, "rectification_deadline_days": 30}
        for i in range(ncrs_pending)
    ]

    record = {
        "project_id":               agreement_no or project_name,
        "contract_id":              agreement_no or project_name,
        "project_name":             project_name,
        "contractor_name":          contractor,
        "reporting_period":         period_end[:7] if period_end else None,
        "reporting_period_end":     period_end,
        "day_number":               day_number or 0,
        "report_date":              period_end,
        "planned_physical_pct":     planned_physical_pct,
        "actual_physical_pct":      actual_physical_pct,
        "prev_physical_pct":        prev_actual_pct,
        "days_since_last_report":   working_days or 30,
        "variance_pct":             variance_pct,
        "cumulative_expenditure_inr": cumulative_exp_inr,
        "financial_progress_pct":   financial_progress_pct,
        "boq_items":                boq_items,
        "materials":                materials,
        "labour_skilled_planned":   labour_skilled_planned,
        "labour_skilled_actual":    labour_skilled_actual,
        "labour_unskilled_planned": labour_unskilled_planned,
        "labour_unskilled_actual":  labour_unskilled_actual,
        "labour_skilled_utilisation_pct":    round(skilled_util, 1),
        "labour_unskilled_utilisation_pct":  round(unskilled_util, 1),
        "labour_deployment_pct":             round(skilled_util, 1),
        "machinery_idle_days":               machinery_idle_days,
        "machinery_deployment_pct":          machinery_deployment_pct,
        "qa_results":               qa_results,
        "test_fail_rate_pct":       round(test_fail_rate, 2),
        "ncrs_pending":             ncrs_pending,
        "open_ncrs":                open_ncrs,
        "rfis_pending":             rfis_pending,
        "days_lost_rainfall":       days_lost_rainfall,
        "days_lost_rainfall_cumulative": days_lost_rainfall,
        "rainfall_mm_monthly":      cumulative_rainfall_mm,
        "weather_anomaly_score":    min(1.0, cumulative_rainfall_mm / 300),
        "row_handover_pct":         round(row_handover_pct, 1),
        "row_pending_km":           row_pending_km,
        "utility_shifting_pending": utility_shifting_pending,
        "railway_clearance_pending": False,
        "forest_clearance_pending": forest_clearance_pending,
        "gfc_drawings_pending":     gfc_pending,
        "ra_bill_number":                  ra_bill_no,
        "ra_bill_amount_inr":              ra_bill_amount_inr,
        "ra_bill_submitted_date":          ra_bill_submitted,
        "payment_received":                prev_payment_yn.strip().upper() == "Y",
        "payment_received_date":           prev_payment_date,
        "payment_delay_days":              payment_delay_days,
        "payment_delayed_streak":          1 if payment_delayed else 0,
        "ld_accumulated_inr":       ld_deducted,
        "performance_security_submitted": True,
        "hindrance_register_unsigned_entries": 0,
        "hindrances":               [],
        "force_majeure_events":     [],
        "variation_orders":         [],
        "ra_bills":                 [],
    }

    errors = validate_mpr(record, prev_actual_pct, bypass_date_check=bypass_date_check)
    if errors:
        raise MPRValidationError(errors)

    return record

def parse_mpr(md_content: str, prev_actual_pct: float = 0.0, bypass_date_check: bool = False) -> dict:
    """
    Parse MPR markdown content into typed exec_data dict.
    Raises MPRValidationError if validation fails.

    Args:
        md_content: raw string content of the .md file
        prev_actual_pct: previous month's actual physical progress % for monotonicity check

    Returns:
        exec_data dict compatible with compliance engine + risk predictor
    """
    text = md_content

    # ── Section 1 — Project metadata ────────────────────────────────────
    project_name   = _kv(text, "Project Name")
    agreement_no   = _kv(text, "Agreement Number")
    contractor     = _kv(text, "Contractor Name")
    eic            = _kv(text, "Engineer-in-Charge")
    period_raw     = _kv(text, "Reporting Period")
    scd_raw        = _kv(text, "Stipulated Date of Completion")
    eot_date       = _kv(text, "Approved EoT Date")
    day_number     = _kv(text, "Day Number", int)

    period_end = None
    if period_raw and " to " in period_raw:
        period_end = period_raw.split(" to ")[1].strip()
        # Normalise to YYYY-MM-DD
        for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                period_end = datetime.strptime(period_end, fmt).strftime("%Y-%m-%d")
                break
            except ValueError:
                continue

    # ── Section 2 — Physical & financial progress ───────────────────────
    planned_physical_pct    = _kv(text, "Planned Physical Progress to Date", float)
    actual_physical_pct     = _kv(text, "Actual Physical Progress to Date", float)
    variance_pct            = _kv(text, "Variance", float)
    cumulative_exp_inr      = _kv(text, "Cumulative Expenditure to Date", float)
    planned_exp_inr         = _kv(text, "Planned Expenditure to Date", float)
    financial_progress_pct  = _kv(text, "Financial Progress", float)

    # ── Section 3 — BoQ table ────────────────────────────────────────────
    boq_items = _table(text, "BoQ Execution")
    if not boq_items:
        boq_items = _table(text, "Section 3")

    # ── Section 4 — Material reconciliation ─────────────────────────────
    materials = _table(text, "Material Reconciliation")
    if not materials:
        materials = _table(text, "Section 4")

    # ── Section 5 — Labour & machinery ───────────────────────────────────
    labour_skilled_planned   = _kv(text, "Planned Skilled Labour", int)
    labour_skilled_actual    = _kv(text, "Actual Skilled Labour", int)
    labour_unskilled_planned = _kv(text, "Planned Unskilled Labour", int)
    labour_unskilled_actual  = _kv(text, "Actual Unskilled Labour", int)
    machinery_idle_days      = _kv(text, "Machinery Idle Days", int) or 0
    machinery_deploy_raw     = _kv(text, "Machinery Deployment") or _kv(text, "Machinery Deployed")
    machinery_deployment_pct = _safe_float(machinery_deploy_raw) if machinery_deploy_raw else 80.0

    # Utilisation percentages
    skilled_util = (
        (labour_skilled_actual / labour_skilled_planned * 100)
        if labour_skilled_planned and labour_skilled_actual is not None
        else 100.0
    )
    unskilled_util = (
        (labour_unskilled_actual / labour_unskilled_planned * 100)
        if labour_unskilled_planned and labour_unskilled_actual is not None
        else 100.0
    )

    # ── Section 6 — QA ────────────────────────────────────────────────────
    qa_results   = _table(text, "Quality Assurance")
    if not qa_results:
        qa_results = _table(text, "Section 6")

    ncrs_issued  = _kv(text, "NCRs Issued This Month", int) or 0
    ncrs_pending = _kv(text, "NCRs Pending Closure", int) or 0
    rfis_submitted = _kv(text, "RFIs Submitted", int) or 0
    rfis_approved  = _kv(text, "RFIs Approved", int) or 0
    rfis_pending   = _kv(text, "RFIs Pending", int) or 0

    # Test fail rate
    total_conducted = sum(_safe_int(r.get("Tests Conducted", 0)) for r in qa_results)
    total_failed    = sum(_safe_int(r.get("Tests Failed", 0)) for r in qa_results)
    test_fail_rate  = (total_failed / total_conducted * 100) if total_conducted > 0 else 0.0

    # Open NCR list
    open_ncrs = [
        {"id": f"NCR-{i+1:03d}", "issued_date": period_end, "rectification_deadline_days": 30}
        for i in range(ncrs_pending)
    ]

    # ── Section 7 — Disruptions & weather ────────────────────────────────
    working_days           = _kv(text, "Working Days in Month", int) or 31
    days_lost_rainfall     = _kv(text, "Days Lost to Rainfall", int) or 0
    days_lost_other        = _kv(text, "Days Lost to Other Hindrances", int) or 0
    daily_rainfall_mm      = _kv(text, "Daily Rainfall", float) or 0.0
    cumulative_rainfall_mm = _kv(text, "Cumulative Rainfall this Month", float) or 0.0

    # ── Section 8 — Land & utilities ─────────────────────────────────────
    row_total_km   = _kv(text, "Total RoW Required", float) or 40.0
    row_handover_km = _kv(text, "RoW Handed Over", float) or row_total_km
    row_pending_km  = _kv(text, "RoW Pending", float) or 0.0
    row_handover_pct = (row_handover_km / row_total_km * 100) if row_total_km > 0 else 100.0

    utility_raw = _kv(text, "Utility Shifting Status") or ""
    utility_shifting_pending = bool(utility_raw and "pending" in utility_raw.lower())
    tree_felling_raw = _kv(text, "Tree Felling Clearance") or "Y"
    forest_clearance_pending = tree_felling_raw.strip().upper() == "N"

    # ── Section 9 — GFC drawings ─────────────────────────────────────────
    gfc_total    = _kv(text, "Total GFC Drawings Required", int) or 0
    gfc_approved = _kv(text, "GFC Drawings Approved", int) or 0
    gfc_pending  = _kv(text, "GFC Drawings Pending", int) or 0

    # ── Section 10 — RA bill & payment ───────────────────────────────────
    ra_bill_no         = _kv(text, "RA Bill Number")
    ra_bill_amount_inr = _kv(text, "RA Bill Amount", float) or 0.0
    ra_bill_submitted  = _kv(text, "RA Bill Submitted Date") or period_end
    prev_payment_yn    = _kv(text, "Previous Bill Payment Received") or "Y"
    prev_payment_date  = _kv(text, "Previous Bill Payment Date")
    payment_delay_days = _kv(text, "Payment Delay", int) or 0

    payment_delayed = payment_delay_days > 0

    ld_deducted_raw = _kv(text, "Cumulative LD Deducted", float) or _kv(text, "LD Deducted", float) or 0.0
    if ld_deducted_raw > 0 and ld_deducted_raw < 10000:
        ld_deducted_raw = ld_deducted_raw * 100000

    # ── Assemble exec_data ────────────────────────────────────────────────
    record = {
        # Metadata
        "project_id":               agreement_no or project_name,
        "contract_id":              agreement_no or project_name,
        "project_name":             project_name,
        "contractor_name":          contractor,
        "reporting_period":         period_end[:7] if period_end else None,  # YYYY-MM
        "reporting_period_end":     period_end,
        "day_number":               day_number or 0,
        "report_date":              period_end,

        # Progress
        "planned_physical_pct":     planned_physical_pct or 0.0,
        "actual_physical_pct":      actual_physical_pct or 0.0,
        "prev_physical_pct":        prev_actual_pct,
        "days_since_last_report":   working_days or 30,
        "variance_pct":             variance_pct or 0.0,
        "cumulative_expenditure_inr": cumulative_exp_inr or 0.0,
        "financial_progress_pct":   financial_progress_pct or 0.0,

        # BoQ + materials
        "boq_items":                boq_items,
        "materials":                materials,

        # Labour
        "labour_skilled_planned":            labour_skilled_planned or 0,
        "labour_skilled_actual":             labour_skilled_actual or 0,
        "labour_unskilled_planned":          labour_unskilled_planned or 0,
        "labour_unskilled_actual":           labour_unskilled_actual or 0,
        "labour_skilled_utilisation_pct":    round(skilled_util, 1),
        "labour_unskilled_utilisation_pct":  round(unskilled_util, 1),
        "labour_deployment_pct":             round(skilled_util, 1),
        "machinery_idle_days":               machinery_idle_days,
        "machinery_deployment_pct":          machinery_deployment_pct,

        # QA
        "qa_results":               qa_results,
        "test_fail_rate_pct":       round(test_fail_rate, 2),
        "ncrs_pending":             ncrs_pending,
        "open_ncrs":                open_ncrs,
        "rfis_pending":             rfis_pending,

        # Weather
        "days_lost_rainfall":       days_lost_rainfall,
        "days_lost_rainfall_cumulative": days_lost_rainfall,  # single-month estimate
        "rainfall_mm_monthly":      cumulative_rainfall_mm,
        "weather_anomaly_score":    min(1.0, cumulative_rainfall_mm / 300),  # rough heuristic

        # Land & utilities
        "row_handover_pct":         round(row_handover_pct, 1),
        "row_pending_km":           row_pending_km,
        "utility_shifting_pending": utility_shifting_pending,
        "railway_clearance_pending": False,
        "forest_clearance_pending": forest_clearance_pending,

        # GFC drawings
        "gfc_drawings_pending":     gfc_pending,

        # RA bill & payment
        "ra_bill_number":                  ra_bill_no,
        "ra_bill_amount_inr":              ra_bill_amount_inr,
        "ra_bill_submitted_date":          ra_bill_submitted,
        "payment_received":                prev_payment_yn.strip().upper() == "Y",
        "payment_received_date":           prev_payment_date,
        "payment_delay_days":              payment_delay_days,
        "payment_delayed_streak":          1 if payment_delayed else 0,

        # Defaults for compliance engine
        "ld_accumulated_inr":       ld_deducted_raw,
        "performance_security_submitted": True,
        "hindrance_register_unsigned_entries": 0,
        "hindrances":               [],
        "force_majeure_events":     [],
        "variation_orders":         [],
        "ra_bills":                 [],
    }

    # ── Validate ──────────────────────────────────────────────────────────
    errors = validate_mpr(record, prev_actual_pct, bypass_date_check=bypass_date_check)
    if errors:
        raise MPRValidationError(errors)

    return record


# ── CLI test ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys, json
    sys.stdout.reconfigure(encoding="utf-8")

    sample = """
# Monthly Progress Report
## Section 1 — Project Metadata
- **Project Name:** NH-44 Karnataka Road Widening (Km 220–260)
- **Agreement Number:** NHAI/KA/EPC/2025/001
- **Contractor Name:** XYZ Constructions Pvt. Ltd.
- **Engineer-in-Charge:** Mr. Ramesh Kumar
- **Reporting Period:** 2025-05-01 to 2025-05-31
- **Stipulated Date of Completion:** 2027-03-31
- **Day Number (from Appointed Date):** 61

## Section 2 — Physical & Financial Progress
- **Planned Physical Progress to Date (%):** 8.2
- **Actual Physical Progress to Date (%):** 6.1
- **Variance (%):** -2.1
- **Cumulative Expenditure to Date (₹):** 11500000
- **Planned Expenditure to Date (₹):** 14200000
- **Financial Progress (%):** 4.6

## Section 5 — Labour & Machinery Deployment
- **Planned Skilled Labour (daily avg):** 85
- **Actual Skilled Labour (daily avg):** 62
- **Planned Unskilled Labour (daily avg):** 210
- **Actual Unskilled Labour (daily avg):** 148
- **Machinery Idle Days (if any):** 0

## Section 6 — Quality Assurance
| Test Type | Material | Tests Conducted | Tests Passed | Tests Failed | Remarks |
|---|---|---|---|---|---|
| Cube Strength (7-day) | M30 Concrete | 6 | 5 | 1 | 1 batch below 20 MPa |
| Field Density Test | Embankment Soil | 4 | 4 | 0 | All ≥97% MDD |

- **NCRs Issued This Month:** 1
- **NCRs Pending Closure:** 1
- **RFIs Submitted:** 3
- **RFIs Approved:** 2
- **RFIs Pending:** 1

## Section 7 — External Disruptions & Hindrance Data
- **Working Days in Month:** 31
- **Days Lost to Rainfall:** 4
- **Cumulative Rainfall this Month (mm):** 48.2

## Section 8 — Land Acquisition & Utilities Status
- **Total RoW Required (km):** 40
- **RoW Handed Over (km):** 36.8
- **RoW Pending (km):** 3.2
- **Utility Shifting Status:** 2 HT lines pending — BESCOM advised 45-day timeline
- **Tree Felling Clearance (Y/N):** Y

## Section 9 — GFC Drawing Status
- **Total GFC Drawings Required:** 48
- **GFC Drawings Approved:** 31
- **GFC Drawings Pending:** 17

## Section 10 — RA Bill & Payment Status
- **RA Bill Number:** RA-02
- **RA Bill Amount (₹):** 5850000
- **RA Bill Submitted Date:** 2025-06-05
- **Previous Bill Payment Received (Y/N):** Y
- **Previous Bill Payment Date:** 2025-04-28
- **Payment Delay (days, if any):** 0
"""

    try:
        result = parse_mpr(sample, prev_actual_pct=3.5)
        print("✅ MPR parsed successfully")
        print(f"   Day: {result['day_number']} | Progress: {result['actual_physical_pct']}%")
        print(f"   Labour skilled util: {result['labour_skilled_utilisation_pct']}%")
        print(f"   QA fail rate: {result['test_fail_rate_pct']}%")
        print(f"   GFC pending: {result['gfc_drawings_pending']}")
        print(f"   RoW: {result['row_handover_pct']}%")
        print(f"   Utility pending: {result['utility_shifting_pending']}")
    except MPRValidationError as e:
        print(f"❌ Validation errors: {e.errors}")

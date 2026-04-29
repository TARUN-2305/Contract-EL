"""
Contract Parser Agent — Deterministic Extraction + Groq LLM fallback
Extracts contract rules from PDF/DOCX via chunking, regex extraction, with
Groq LLM as fallback for ambiguous fields.
Produces: rule_store_{contract_id}.json
"""
import gc
import json
import os
import re
from sentence_transformers import SentenceTransformer
from sqlalchemy.orm import Session

from db.database import SessionLocal
from db.models import RuleStore
from db.vector_store import VectorStore
from agents.extraction_engine import deterministic_extract
from utils.groq_client import groq_json_extract

# ── Embedding model (loaded once) ──────────────────────────────────────
_embed_model = None

def get_embed_model():
    global _embed_model
    if _embed_model is None:
        print("[ParserAgent] Loading ST model (all-MiniLM-L6-v2)...")
        _embed_model = SentenceTransformer("all-MiniLM-L6-v2")
        print("[ParserAgent] Model loaded.")
    return _embed_model

# ── PDF Text Extraction ────────────────────────────────────────────────

def extract_text_from_pdf(pdf_path: str) -> list[dict]:
    """Extract text from PDF, one entry per page."""
    pages = []
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages.append({"page_number": i + 1, "text": text})
    except Exception as e:
        print(f"[ParserAgent] pdfplumber failed, falling back to pypdf: {e}")
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page_number": i + 1, "text": text})
    return pages

def extract_text_from_docx(docx_path: str) -> list[dict]:
    """Extract text from DOCX preserving structure via docx_to_md converter."""
    try:
        from utils.docx_to_md import docx_to_md
        md_content = docx_to_md(docx_path)
        return [{"page_number": 1, "text": md_content}]
    except Exception as e:
        print(f"[ParserAgent] DOCX→MD conversion failed: {e}, falling back to raw extract")
        pages = []
        try:
            import docx
            doc = docx.Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    full_text.append(row_text)
            content = "\n\n".join(full_text)
            pages.append({"page_number": 1, "text": content})
        except Exception as fallback_e:
            print(f"[ParserAgent] Fallback DOCX extraction failed: {fallback_e}")
        return pages


# ── Semantic Chunking ──────────────────────────────────────────────────

ARTICLE_PATTERN = re.compile(
    r"(?:^|\n)\s*(?:ARTICLE|Article|CLAUSE|Clause|SECTION|Section)\s+(\d+[\.\d]*)",
    re.MULTILINE,
)

def chunk_contract_text(pages: list[dict]) -> list[dict]:
    """
    Split extracted pages into semantic chunks respecting article/clause boundaries.
    Each chunk is tagged with {clause_id, page_number, section_type, text}.
    """
    full_text = "\n\n".join(p["text"] for p in pages)
    page_map = {}
    offset = 0
    for p in pages:
        page_map[offset] = p["page_number"]
        offset += len(p["text"]) + 2  # +2 for \n\n

    # Find all article/clause boundaries
    splits = list(ARTICLE_PATTERN.finditer(full_text))

    chunks = []
    if not splits:
        # No structured headings found — chunk by paragraphs
        paragraphs = full_text.split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if len(para) > 50:
                chunks.append({
                    "clause_id": None,
                    "section_type": "paragraph",
                    "page_number": 1,
                    "text": para,
                })
        return chunks

    # First chunk: everything before the first article
    preamble = full_text[: splits[0].start()].strip()
    if len(preamble) > 50:
        chunks.append({
            "clause_id": "PREAMBLE",
            "section_type": "preamble",
            "page_number": 1,
            "text": preamble,
        })

    for i, match in enumerate(splits):
        start = match.start()
        end = splits[i + 1].start() if i + 1 < len(splits) else len(full_text)
        clause_text = full_text[start:end].strip()
        clause_id = match.group(1)

        # Determine page number
        page_num = 1
        for off, pn in sorted(page_map.items()):
            if start >= off:
                page_num = pn

        # Determine section type from clause_id
        section_type = "article" if "." not in clause_id else "sub_clause"

        chunks.append({
            "clause_id": clause_id,
            "section_type": section_type,
            "page_number": page_num,
            "text": clause_text,
        })

    return chunks


# ── Hierarchical Extraction Plan ───────────────────────────────────────

EXTRACTION_PLAN = [
    {"target": "milestones",           "query": "project milestone schedule completion date"},
    {"target": "liquidated_damages",   "query": "liquidated damages compensation delay penalty rate"},
    {"target": "performance_security", "query": "performance security guarantee bank deposit"},
    {"target": "quality_assurance",    "query": "quality assurance testing NCR non-conformance"},
    {"target": "force_majeure",        "query": "force majeure acts of god political event notice"},
    {"target": "eot_rules",            "query": "extension of time hindrance delay application"},
    {"target": "variation_orders",     "query": "change of scope variation order extra items"},
    {"target": "termination",          "query": "termination default notice cure period"},
    {"target": "dispute_resolution",   "query": "dispute arbitration DRB conciliation escalation"},
    {"target": "bonus",                "query": "early completion incentive bonus clause 2A"},
    {"target": "conditions_precedent", "query": "conditions precedent appointed date handover"},
    {"target": "payment_workflow",     "query": "running account bill measurement certificate payment"},
]

EXTRACTION_PROMPTS = {
    "milestones": """You are a legal extraction engine for Indian infrastructure contracts.
Extract ALL milestone definitions from the following contract text.
Return ONLY valid JSON. No preamble, no explanation, no markdown.

Required fields per milestone:
- id (M1, M2, etc.)
- name
- trigger_pct_of_scp (% of scheduled construction period, integer)
- trigger_day (absolute day number, integer)
- required_physical_progress_pct (integer)
- ld_rate_pct_per_day (float)
- ld_basis ("apportioned_milestone_value" or "total_contract_price")
- catch_up_refund_eligible (boolean)
- source_clause (string, exact article/clause number)

If a field is not found, use null. Never hallucinate values.

Contract text:
{context}""",

    "liquidated_damages": """Extract liquidated damages rules from this contract text.
Return ONLY valid JSON with these fields:
- daily_rate_pct (float)
- max_cap_pct (int)
- max_cap_inr (float, calculate from contract value if given)
- catch_up_refund (boolean)
- source_clause (string)

Contract text:
{context}""",

    "performance_security": """Extract performance security requirements from this contract text.
Return ONLY valid JSON with these fields:
- pct_of_contract_value (float)
- amount_inr (float)
- acceptable_forms (list of strings)
- submission_deadline_days (int)
- late_fee_pct_per_day (float)
- max_extension_days (int)
- consequence_of_failure (string)
- source_clause (string)

Contract text:
{context}""",

    "force_majeure": """Extract force majeure rules from this contract text.
Return ONLY valid JSON with these fields:
- notice_deadline_days (int)
- notice_recipient (string)
- required_notice_contents (list of strings)
- ongoing_reporting_frequency (string)
- max_suspension_days_before_termination (int)
- categories (list of strings)
- source_clause (string)

Contract text:
{context}""",

    "eot_rules": """Extract Extension of Time rules from this contract text.
Return ONLY valid JSON with these fields:
- application_deadline_days (int)
- hindrance_register_mandatory (boolean)
- overlap_deduction_required (boolean)
- source_clause (string)

Contract text:
{context}""",

    "termination": """Extract termination rules from this contract text.
Return ONLY valid JSON with these fields:
- contractor_default_triggers (list of objects with trigger, threshold_days, cure_period_days)
- source_clause (string)

Contract text:
{context}""",

    "dispute_resolution": """Extract dispute resolution rules from this contract text.
Return ONLY valid JSON with these fields:
- tiers (list of objects with tier, mechanism, deadline_days)
- source_clause (string)

Contract text:
{context}""",

    "bonus": """Extract early completion bonus rules from this contract text.
Return ONLY valid JSON with these fields:
- applicable (boolean)
- rate_pct_per_month (float)
- max_cap_pct (float)
- source_clause (string)

Contract text:
{context}""",

    "payment_workflow": """Extract payment workflow rules from this contract text.
Return ONLY valid JSON with these fields:
- ra_bill_submission_day (int)
- verification_deadline_days (int)
- payment_release_deadline_days (int)
- mandatory_deductions (list of objects with type, rate_pct)
- source_clause (string)

Contract text:
{context}""",
}

# Default prompt for targets without a specific template
DEFAULT_EXTRACTION_PROMPT = """Extract {target} rules from this contract text.
Return ONLY valid JSON. No preamble, no explanation.
If a field is not found, use null. Never hallucinate values.

Contract text:
{context}"""


# ── Validation Rules ───────────────────────────────────────────────────

VALIDATION_RULES = {
    "milestones": {
        "ld_rate_pct_per_day": lambda x: x is None or (0 < x <= 1),
        "trigger_pct_of_scp": lambda x: x is None or (0 < x <= 100),
    },
    "performance_security": {
        "pct_of_contract_value": lambda x: x is None or (4 <= x <= 10),
        "submission_deadline_days": lambda x: x is None or x == 15,
    },
    "force_majeure": {
        "notice_deadline_days": lambda x: x is None or x == 7,
    },
}

def validate_extracted(target: str, data) -> list[str]:
    """Validate extracted data against rules. Returns list of warnings."""
    warnings = []
    rules = VALIDATION_RULES.get(target, {})

    if isinstance(data, list):
        for i, item in enumerate(data):
            if isinstance(item, dict):
                for field, validator in rules.items():
                    val = item.get(field)
                    if val is not None and not validator(val):
                        warnings.append(f"{target}[{i}].{field}={val} failed validation")
    elif isinstance(data, dict):
        for field, validator in rules.items():
            val = data.get(field)
            if val is not None and not validator(val):
                warnings.append(f"{target}.{field}={val} failed validation")

    return warnings


# ── Parser Agent Class ─────────────────────────────────────────────────

class ParserAgent:
    def __init__(self):
        """Parser agent using deterministic regex extraction + Groq LLM fallback."""
        self.vector_store = VectorStore()

    def parse_contract(
        self,
        file_path: str,
        contract_id: str,
        contract_type: str,
        contract_value_inr: float,
        scp_days: int,
        project_name: str,
        location: str,
        contractor_name: str,
    ) -> dict:
        """Full parsing pipeline: File (PDF/DOCX) → chunks → embeddings → extraction → rule store."""
        print(f"[ParserAgent] Starting parse for contract {contract_id}")

        # Stage 1: Extract text
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            print(f"[ParserAgent] Stage 1: Extracting text from DOCX ({file_path})...")
            pages = extract_text_from_docx(file_path)
        else:
            print(f"[ParserAgent] Stage 1: Extracting text from PDF ({file_path})...")
            pages = extract_text_from_pdf(file_path)
        print(f"[ParserAgent] Extracted {len(pages)} pages")

        # Stage 2: Semantic chunking
        print("[ParserAgent] Stage 2: Chunking contract text...")
        chunks = chunk_contract_text(pages)
        print(f"[ParserAgent] Created {len(chunks)} chunks")

        # Stage 3: Embed and store
        print("[ParserAgent] Stage 3: Embedding chunks...")
        model = get_embed_model()
        texts = [c["text"] for c in chunks]
        embeddings = model.encode(texts).tolist()

        db = SessionLocal()
        try:
            self.vector_store.store_chunks(db, contract_id, chunks, embeddings)
        finally:
            db.close()

        # Free embedding model from memory before extraction
        global _embed_model
        _embed_model = None
        gc.collect()
        print("[ParserAgent] Embedding model released from memory.")

        # Stage 4: Deterministic extraction
        print("[ParserAgent] Stage 4: Deterministic extraction...")
        full_text = "\n\n".join(p["text"] for p in pages)
        extracted = deterministic_extract(full_text)
        audit_log = {target: {"method": "deterministic_regex", "warnings": []} for target in extracted}
        unresolved = {t: ["null result"] for t, v in extracted.items() if v is None}

        # Validate key fields
        for target, data in extracted.items():
            warnings = validate_extracted(target, data)
            if warnings:
                audit_log[target]["warnings"] = warnings
                unresolved.setdefault(target, []).extend(warnings)

        # Stage 5: Assemble rule store
        print("[ParserAgent] Stage 5: Assembling rule store...")
        from datetime import date
        rule_store = {
            "contract_id": contract_id,
            "contract_type": contract_type,
            "contract_value_inr": contract_value_inr,
            "scp_days": scp_days,
            "project_name": project_name,
            "location": location,
            "contractor_name": contractor_name,
            "appointed_date": None,
            "scheduled_completion_date": None,
        }
        rule_store.update(extracted)

        # Patch M4 trigger_day to scp_days if extraction returned None
        milestones = rule_store.get("milestones") or []
        for m in milestones:
            if m.get("id") == "M4" and not m.get("trigger_day"):
                m["trigger_day"] = scp_days
                print(f"[ParserAgent] M4 trigger_day patched to scp_days={scp_days}")


        # Write to filesystem
        os.makedirs("data/rule_store", exist_ok=True)
        os.makedirs("data/audit", exist_ok=True)

        rule_store_path = f"data/rule_store/rule_store_{contract_id}.json"
        with open(rule_store_path, "w", encoding="utf-8") as f:
            json.dump(rule_store, f, indent=2, default=str)
        print(f"[ParserAgent] Rule store written to {rule_store_path}")

        audit_path = f"data/audit/extraction_audit_{contract_id}.json"
        with open(audit_path, "w", encoding="utf-8") as f:
            json.dump(audit_log, f, indent=2)

        if unresolved:
            unresolved_path = f"data/audit/unresolved_fields_{contract_id}.json"
            with open(unresolved_path, "w", encoding="utf-8") as f:
                json.dump(unresolved, f, indent=2)
            print(f"[ParserAgent] [WARN] Unresolved fields written to {unresolved_path}")

        # Write to database
        db = SessionLocal()
        try:
            db_rule = RuleStore(project_id=contract_id, rules=rule_store)
            db.add(db_rule)
            db.commit()
            print(f"[ParserAgent] Rule store saved to database")
        finally:
            db.close()

        print(f"[ParserAgent] [OK] Contract parsing complete for {contract_id}")
        return rule_store
